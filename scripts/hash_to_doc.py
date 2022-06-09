'''
https://pythoner.name/walk
https://tokmakov.msk.ru/blog/item/78
https://docs-python.ru/packages/modul-python-docx-python/
https://vc.ru/dev/185015-analiz-dokumentov-word-s-ispolzovaniem-python
https://docs-python.ru/packages/modul-python-docx-python/obekt-table/
https://ru.stackoverflow.com/questions/1332992/%D0%90%D0%B2%D1%82%D0%BE%D0%BC%D0%B0%D1%82%D0%B8%D0%B7%D0%B0%D1%86%D0%B8%D1%8F-%D1%81%D0%BE%D0%B7%D0%B4%D0%B0%D0%BD%D0%B8%D1%8F-%D1%82%D0%B0%D0%B1%D0%BB%D0%B8%D1%86-python-docx

'''

import docx
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
# from re import match
from os import getcwd, listdir, walk, path


def get_hash_files(dir_path):
    file_list = []
    for address, dirs, files in walk(dir_path):
        for name in files:
            if name.endswith('.py'):
                file_list.append(name)
            # print(path.join(address, name))
    return file_list


def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)
    table = document.add_table(rows=1, cols=cols_number)
    # table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):  # Заполнение шапки таблицы
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])
    return table


def doc_build(docx_file, files_path=''):
    if files_path == '':
        files_path = getcwd()
    if docx_file == '':
        docx_file = "Таблица хешей.docx"

    computer = 'Comp'
    is_dir = 'Каталог'  # что за ресурс прохеширован
    '''
    ======= Начинаем наполнение файла docx. ========
    '''
    document = docx.Document(docx=docx_file)

    # properties = document.core_properties
    # print('Автор документа:', properties.author)
    # print('Автор последней правки:', properties.last_modified_by)
    # print('Дата создания документа:', properties.created)
    # print('Дата последней правки:', properties.modified)
    # print('Дата последней печати:', properties.last_printed)
    # print('Количество сохранений:', properties.revision)

    hash_files = get_hash_files(files_path)

    rows2 = [('', f, 'files_path', 'HASH') for f in hash_files]
    print(rows2)
    headers = ('Наменование ТС:', computer, '', '')
    records_table0 = (
        ('Алгоритм контрольной суммы:', 'SHA1', 'Nan', 'Nan'),
        ('Наименование ИС', is_dir, 'Nan', 'Nan'),
        ('№', 'Имя файла', 'Размещение файла', 'Хэш сумма'),
        # ('', '', '', ''),
        *rows2
    )
    table0 = create_table(document, headers, records_table0)

    # ======== Наполнение таблицы, добавляя новые строки. =============
    # for n, desc in enumerate(params_list):
    #     row_cells = table.add_row().cells
    #     run01 = row_cells[0].paragraphs[0].add_run(str(desc))
    #     # run01.bold = True
    #     run02 = row_cells[1].paragraphs[0].add_run('\n'.join(table_dict.get(n)))
    #     # row_cells[0].text = str(desc)
    #     # row_cells[1].text = '\n'.join(table_dict.get(n))

    document.add_paragraph()

    # Тут пример цикла таблицы
    # rows = [
    #     [x, x, x * x] for x in range(1, 10)
    # ]
    # table1 = create_table(document, ('x', 'y', 'x * y'), rows)

    # document.add_page_break()  # Добавление разрыва страницы
    document.save(f'{docx_file}')
    print(f'Сохранено в "{docx_file}"')


if __name__ == '__main__':
    doc_build('Про хэширование.docx')
