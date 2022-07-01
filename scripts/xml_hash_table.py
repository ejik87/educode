import docx
from docx.shared import Inches, Cm, Pt
# from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree
from os import getcwd, path, walk


def parseXML(xmlFile):
    """Parse the XML file"""
    # parser = etree.XMLParser(remove_blank_text=True)  # lxml.etree only!
    file_hash_dict = dict()
    file_types = (
        # 'DLL',
        'RTCX.EXE',
        'USER.TMSD',
        'USER_16.TMSD',
        '.UP'
        '.RTP',
        '.CNV',
        '.DBB',
        'MPS.EXE',
        'PCOMM.DLL',
        'INSATHDASVR.DLL',
        'VIEWX.EXE',
        'VIEWXCORE.DLL',
        'SCXCMD.EXE',
        'SERVERCONFIG.EXE',
        'DBCLIENT.DLL',
        'DBKERNEL.DLL',
        'OBJECTS.CL',
        'OELITE.CFG',
        'SQLC.CFG',
        'WINCCEXPLORER.EXE',
        'WINCCCONFIGURATOR.EXE',
        'DB.DLL',
        'WINCC50.DLL',
        'WINCC50X64.DLL',
        'АРМ УЗРГ.exe',
        'АРМ УЗРГ.vshost.exe'
    )

    tree = etree.parse(xmlFile)  # Читаем из XML файла инфу
    file_list = tree.findall('.//name')
    hash_list = tree.findall('.//SHA1')

    for f, h in zip(file_list, hash_list):
        file_name = f.text.split('\\')[-1]  # Отделяем имя файла от пути в тэге name
        # name_file.endswith('EXE')
        # fff = file_name.split('.')[-1].upper()
        # if fff in file_types:
        if file_name.upper().endswith(file_types):
            file_hash_dict[file_name] = [f.text.replace(file_name, ''), h.text] # Добавляем -> 'файл': ['SHA1 контрольную сумму', путь к файлу]

    return file_hash_dict


# def doc_file():
#     files_list = []
#     path_folder = getcwd()
#     for root, dirs, files in walk(path_folder):
#         for file in files:
#             if file.endswith('docx') and not file.startswith('~'):
#                 files_list.append(path.join(root, file))
#     return files_list


# def show_doc(paths):
#     for p in paths:
#         doc = docx.Document(p)
#         properties = doc.core_properties
#         print('Автор документа:', properties.author)
#         print('Автор последней правки:', properties.last_modified_by)
#         print('Дата создания документа:', properties.created)
#         print('Дата последней правки:', properties.modified)
#         print('Дата последней печати:', properties.last_printed)
#         print('Количество сохранений:', properties.revision)
#
#         for table in doc.tables:
#             for index, row in enumerate(table.rows):
#                 if index == 0:
#                     row_text = list(cell.text for cell in row.cells)
#                     if 'Имя файла' not in row_text or 'Значение контрольной суммы' not in row_text:
#                         break
#                 for cell in row.cells:
#                     print(cell.text)
#
#         for table in doc.tables:
#             unique, merged = set(), set()
#             for row in table.rows:
#                 for cell in row.cells:
#                     tc = cell._tc
#                     cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
#                     if cell_loc in unique:
#                         merged.add(cell_loc)
#                     else:
#                         unique.add(cell_loc)
#             print(merged)


def get_hash_files(dir_path):
    file_list = []
    for address, dirs, files in walk(dir_path):
        for name in files:
            if name.endswith('.xml') and name.startswith(('ia', 'po')):
                file_list.append(path.join(address, name))
    return file_list


def parse_hosts(file_path):
    names = [file.split('/')[-1] for file in file_path]
    hosts = []
    for host in names:
        if host.endswith('.xml') and host.startswith(('ia', 'po')):
            host = host.rstrip('.xml')
            name = host[host.find('hashes-') + 7:]
            if name not in hosts:
                hosts.append(name)
    return hosts


# def create_table(document, headers, rows, style='Table Grid'):
#     cols_number = len(headers[0])
#     table = document.add_table(rows=1, cols=cols_number)
#     table.style = style
#
#     table.cell(0, 0).width = Cm(1.0)  # Настройка ширины 1го столбца
#
#     hdr_cells = table.rows[0].cells
#     # for i in range(cols_number):
#     #     hdr_cells[i].text = headers[i]
#     for row in headers:  # Заполнение шапки таблицы
#         if row is not headers[0]:  # если строка шапки не первая
#             hdr_cells = table.add_row().cells
#         for i in range(cols_number):
#             hdr_cells[i].text = str(row[i])
#
#     for row in rows:
#         row_cells = table.add_row().cells
#         for i in range(cols_number):
#             row_cells[i].text = str(row[i])
#     return table


def doc_build(docx_file='', files_path=getcwd()):
    if docx_file == '':
        document = docx.Document()
        docx_file = "Таблица хешей.docx"
    document = docx.Document(docx=docx_file)  # Открываем файл для добавления таблиц
    '''
    ======= Начинаем наполнение файла docx. ========
    '''
    # создаем пустую таблицу
    # если сразу создать таблицу, ее ширина потом не устанавливается из-за бага:
    # https://github.com/python-openxml/python-docx/issues/315
    # table00 = document.add_table(rows=0, cols=0, style='Table Grid')
    # # прижимаем таблицу вправо
    # table00.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # # добавляем колонку, строку
    # column00 = table00.add_column(Cm(1.5))
    # row00 = table00.add_row()
    # # заполняем ячейку текстом
    # cells = row00.cells
    # cells[0].text = "Lorem ipsum dolor"

    style_doc = document.styles['Normal']
    font_p = style_doc.font
    font_p.name = 'Times New Roman'
    font_p.size = Pt(12)  # 12900 * 12  # Устанавливаем шрифт 12
    font_p.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)

    hash_files = get_hash_files(files_path)  # Ищем все файлы XML
    hostnames = parse_hosts(hash_files)  # Выделяем hostname из имём файлов с хэшами

    type_dir = {  # что за ресурс прохеширован
        'ia': 'Инф. актив',
        'po': 'Прикладное ПО',
    }

    for hostname in hostnames:  # Цикл создания таблиц
        ''' Фильтруем имена хэш-файлов по имени хоста'''
        file_from_hostname = [filename for filename in hash_files if hostname == filename[filename.find('hashes-') + 7:].rstrip('.xml')]
        rows = []

        header = (
            '№',
            'Имя файла / идентификатор объекта (если не файл)',
            'Место размещения',
            'Значение контрольной суммы'
        )

        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.cell(0, 0).width = Cm(1.0)  # Настройка ширины 1го столбца

        '''Тут шапка таблиц'''
        hdr_cells = table.rows[0].cells
        # table.cell(0, 0).merge(table.cell(0, 3))
        hdr_cells[0].merge(hdr_cells[-1])  # объединение ячеек
        table.cell(0, 0).text = str(f'Наименование ТС: {hostname}')
        hdr_cells = table.add_row().cells
        # table.cell(1, 0).merge(table.cell(1, 3))
        hdr_cells[0].merge(hdr_cells[-1])  # объединение ячеек
        table.cell(1, 0).text = str(f'Используемый алгоритм подсчёта значений контрольный сумм: SHA1')
        hdr_cells = table.add_row().cells
        for i, row in enumerate(header):
            table.cell(2, i).text = str(row)
        '''Конец шапки'''

        type_flag = ''
        row_number = 1
        for hash_file in file_from_hostname:
            # if hostname == hash_file[hash_file.find('hashes-') + 7:].rstrip('.xml'):
            try:
                hashes_dict = parseXML(hash_file)
                if hashes_dict:
                    type_comp = type_dir[hash_file.split('/')[-1][:2]]  # отделяем символы типа содержимого файла (io, po)
                    if type_comp != type_flag:
                        type_flag = type_comp
                        row_number = 1
                        # Добавляем разделитель-шапку
                        row_cells = table.add_row().cells
                        row_cells[0].merge(row_cells[-1])
                        row_cells[0].text = str(f'Наименование подсистемы / компонента: {type_comp}')

                    '''# Заполняем табличку данными'''
                    # Достаём итемы из словаря, имя файла:[путь, хэш]
                    rows = ([(n, i[0], i[1][0], i[1][1]) for n, i in enumerate(hashes_dict.items(), row_number)])
                    for row in rows:
                        row_cells = table.add_row().cells
                        for i in range(4):
                            row_cells[i].text = str(row[i])
                    row_number += len(rows)

            except OSError:
                print(OSError)
        document.add_paragraph()

    # Тут пример цикла таблицы
    # rows3 = [
    #     [x, x, x * x] for x in range(1, 10)
    # ]
    # table1 = create_table(document, ('x', 'y', 'x * y'), rows3)

    # document.add_page_break()  # Добавление разрыва страницы
    document.save(f'{docx_file}')
    print(f'Сохранено в "{getcwd()}/{docx_file}"')


if __name__ == '__main__':
    doc_build('документ по контролю целостности.docx')
    
